#!/usr/bin/env python3
"""Convert Markdown to DOCX using a Digica template with Outfit font styling.

Usage:
    python md2docx.py input.md [output.docx]

If output is not specified, creates input.docx next to the input file.

Dependencies: python-docx, markdown-it-py, Pillow
"""

import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from markdown_it import MarkdownIt
from markdown_it.token import Token

SCRIPT_DIR = Path(__file__).parent
ASSETS_DIR = SCRIPT_DIR.parent / "assets"
TEMPLATE_PATH = ASSETS_DIR / "template.docx"
FONT_DIR = ASSETS_DIR / "fonts"

FONT_NAME = "Outfit"
HEADING_STYLES = {
    1: {"size": Pt(20), "bold": True, "space_before": Pt(24), "space_after": Pt(8)},
    2: {"size": Pt(16), "bold": True, "space_before": Pt(18), "space_after": Pt(6)},
    3: {"size": Pt(14), "bold": True, "space_before": Pt(14), "space_after": Pt(4)},
    4: {"size": Pt(12), "bold": True, "space_before": Pt(12), "space_after": Pt(4)},
    5: {"size": Pt(11), "bold": True, "space_before": Pt(10), "space_after": Pt(4)},
    6: {"size": Pt(10), "bold": True, "space_before": Pt(10), "space_after": Pt(4)},
}
BODY_SIZE = Pt(11)
CODE_FONT = "Courier New"
CODE_BG_COLOR = "E8E8E8"
MAX_IMAGE_WIDTH = Cm(16)


def create_document() -> Document:
    doc = Document(str(TEMPLATE_PATH))
    return doc


def set_run_font(run, font_name=FONT_NAME, size=BODY_SIZE, bold=None, italic=None):
    run.font.name = font_name
    run.font.size = size
    rpr = run._element.get_or_add_rPr()
    rpr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}" w:eastAsia="{font_name}"/>'))
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_heading(doc: Document, text: str, level: int) -> None:
    style = HEADING_STYLES.get(level, HEADING_STYLES[6])
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    pf = p.paragraph_format
    pf.space_before = style["space_before"]
    pf.space_after = style["space_after"]
    run = p.add_run(text)
    set_run_font(run, size=style["size"], bold=style["bold"])


def add_paragraph_with_inlines(doc: Document, tokens: list[Token], md_dir: Path) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    _render_inlines_to_paragraph(p, tokens, md_dir)


def _render_inlines_to_paragraph(p, tokens: list[Token], md_dir: Path,
                                  base_bold=False, base_italic=False) -> None:
    bold_stack = base_bold
    italic_stack = base_italic
    link_url = None

    i = 0
    while i < len(tokens):
        tok = tokens[i]

        if tok.type == "text":
            run = p.add_run(tok.content)
            set_run_font(run, bold=bold_stack or None, italic=italic_stack or None)
            if link_url:
                _add_hyperlink_style(run)

        elif tok.type == "code_inline":
            run = p.add_run(tok.content)
            set_run_font(run, font_name=CODE_FONT, size=BODY_SIZE)
            _set_run_shading(run, CODE_BG_COLOR)

        elif tok.type == "softbreak":
            run = p.add_run("\n")
            set_run_font(run)

        elif tok.type == "hardbreak":
            run = p.add_run()
            run.add_break()

        elif tok.type == "strong_open":
            bold_stack = True
        elif tok.type == "strong_close":
            bold_stack = base_bold

        elif tok.type == "em_open":
            italic_stack = True
        elif tok.type == "em_close":
            italic_stack = base_italic

        elif tok.type == "link_open":
            link_url = tok.attrGet("href")
        elif tok.type == "link_close":
            if link_url:
                run = p.add_run(f" [{link_url}]")
                set_run_font(run, size=Pt(9), italic=True)
                run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
            link_url = None

        elif tok.type == "image":
            src = tok.attrGet("src")
            if src:
                img_path = _resolve_image_path(src, md_dir)
                if img_path and img_path.exists():
                    _add_centered_image(p, img_path)

        i += 1


def _set_run_shading(run, color: str):
    rpr = run._element.get_or_add_rPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color}"/>')
    rpr.append(shading)


def _add_hyperlink_style(run):
    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
    run.underline = True


def _resolve_image_path(src: str, md_dir: Path) -> Path | None:
    p = Path(src)
    if p.is_absolute():
        return p
    resolved = md_dir / p
    return resolved


def _add_centered_image(paragraph, img_path: Path):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    max_w_inches = 6.3  # ~16cm
    try:
        from PIL import Image as PILImage
        with PILImage.open(img_path) as img:
            w_px, _ = img.size
            dpi_info = img.info.get("dpi", (96, 96))
            dpi_x = dpi_info[0] if dpi_info[0] > 0 else 96
            width_inches = w_px / dpi_x
            if width_inches > max_w_inches:
                width_inches = max_w_inches
            run.add_picture(str(img_path), width=Inches(width_inches))
    except Exception:
        run.add_picture(str(img_path), width=Inches(max_w_inches))


def add_code_block(doc: Document, code: str, language: str = "") -> None:
    p = doc.add_paragraph()
    _set_paragraph_shading(p, CODE_BG_COLOR)
    _set_paragraph_borders(p, "D0D0D0")
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(8)
    pf.left_indent = Cm(0.6)
    pf.right_indent = Cm(0.6)

    lines = code.rstrip("\n").split("\n")
    for i, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, font_name=CODE_FONT, size=Pt(9))
        if i < len(lines) - 1:
            run = p.add_run()
            run.add_break()


def _set_paragraph_shading(paragraph, color: str):
    ppr = paragraph._element.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color}"/>')
    ppr.append(shading)


def _set_paragraph_borders(paragraph, color: str):
    ppr = paragraph._element.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="4" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="4" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="4" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="4" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    ppr.append(borders)


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              alignments: list[str | None] | None = None) -> None:
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)

    # Header row
    for j, header_text in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header_text.strip())
        set_run_font(run, bold=True)
        _set_cell_shading(cell, "D9E2F3")
        if alignments and j < len(alignments):
            _apply_alignment(p, alignments[j])

    # Data rows
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j >= num_cols:
                break
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text.strip())
            set_run_font(run)
            if alignments and j < len(alignments):
                _apply_alignment(p, alignments[j])


def _set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def _set_cell_shading(cell, color: str):
    tc_pr = cell._element.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color}"/>')
    tc_pr.append(shading)


def _apply_alignment(paragraph, align: str | None):
    if align == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_list_item(doc: Document, tokens: list[Token], md_dir: Path,
                  ordered: bool, level: int = 0) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(1.27 * (level + 1))
    pf.first_line_indent = Cm(-0.63)
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)

    if not ordered:
        bullet = "\u2022" if level == 0 else ("\u25E6" if level == 1 else "\u25AA")
        run = p.add_run(f"{bullet} ")
        set_run_font(run)
    # For ordered lists, the number prefix is added by the caller

    _render_inlines_to_paragraph(p, tokens, md_dir)


# --- Markdown parsing and rendering ---

def parse_markdown(md_text: str) -> list[Token]:
    md = MarkdownIt("commonmark", {"breaks": False})
    md.enable("table")
    tokens = md.parse(md_text)
    return tokens


def render_tokens(doc: Document, tokens: list[Token], md_dir: Path,
                  list_stack: list[dict] | None = None) -> None:
    i = 0
    if list_stack is None:
        list_stack = []

    while i < len(tokens):
        tok = tokens[i]

        if tok.type == "heading_open":
            level = int(tok.tag[1])  # h1 -> 1, h2 -> 2, etc.
            inline_tok = tokens[i + 1]  # heading content (inline)
            text = inline_tok.content
            add_heading(doc, text, level)
            i += 3  # heading_open, inline, heading_close
            continue

        if tok.type == "paragraph_open":
            inline_tok = tokens[i + 1]
            if inline_tok.children:
                # Check if it's just an image
                if len(inline_tok.children) == 1 and inline_tok.children[0].type == "image":
                    img_tok = inline_tok.children[0]
                    src = img_tok.attrGet("src")
                    if src:
                        img_path = _resolve_image_path(src, md_dir)
                        if img_path and img_path.exists():
                            p = doc.add_paragraph()
                            _add_centered_image(p, img_path)
                            i += 3
                            continue
                add_paragraph_with_inlines(doc, inline_tok.children, md_dir)
            else:
                p = doc.add_paragraph()
                run = p.add_run(inline_tok.content)
                set_run_font(run)
            i += 3  # paragraph_open, inline, paragraph_close
            continue

        if tok.type == "bullet_list_open":
            list_stack.append({"ordered": False, "counter": 0})
            i += 1
            continue

        if tok.type == "ordered_list_open":
            list_stack.append({"ordered": True, "counter": 0})
            i += 1
            continue

        if tok.type in ("bullet_list_close", "ordered_list_close"):
            if list_stack:
                list_stack.pop()
            i += 1
            continue

        if tok.type == "list_item_open":
            level = len(list_stack) - 1
            current_list = list_stack[-1] if list_stack else {"ordered": False, "counter": 0}
            current_list["counter"] += 1

            # Collect inline content from the list item
            # Next tokens: paragraph_open, inline, paragraph_close, ... , list_item_close
            # But there may be nested lists too
            j = i + 1
            inline_tokens = []
            # Find the inline content of this list item (first paragraph)
            while j < len(tokens) and tokens[j].type != "list_item_close":
                if tokens[j].type == "paragraph_open":
                    inline_tok = tokens[j + 1]
                    if inline_tok.children:
                        inline_tokens = inline_tok.children
                    else:
                        # Create a simple text token
                        inline_tokens = [Token(type="text", tag="", nesting=0, content=inline_tok.content)]
                    j += 3  # skip paragraph_open, inline, paragraph_close
                    break
                j += 1

            # Render list item
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Cm(1.27 * (level + 1))
            pf.first_line_indent = Cm(-0.63)
            pf.space_before = Pt(2)
            pf.space_after = Pt(2)

            if current_list["ordered"]:
                prefix = f"{current_list['counter']}. "
            else:
                bullet = "\u2022" if level == 0 else ("\u25E6" if level == 1 else "\u25AA")
                prefix = f"{bullet} "

            run = p.add_run(prefix)
            set_run_font(run)
            _render_inlines_to_paragraph(p, inline_tokens, md_dir)

            # Now handle nested lists within this list_item
            depth = 1
            while j < len(tokens):
                if tokens[j].type == "list_item_close":
                    depth -= 1
                    if depth == 0:
                        i = j + 1
                        break
                    j += 1
                elif tokens[j].type == "list_item_open":
                    depth += 1
                    j += 1
                elif tokens[j].type in ("bullet_list_open", "ordered_list_open"):
                    # Nested list — render recursively from here
                    nested_end = _find_matching_close(tokens, j)
                    render_tokens(doc, tokens[j:nested_end + 1], md_dir, list_stack)
                    j = nested_end + 1
                else:
                    j += 1
            else:
                i = j
            continue

        if tok.type == "fence":
            add_code_block(doc, tok.content, tok.info)
            i += 1
            continue

        if tok.type == "code_block":
            add_code_block(doc, tok.content)
            i += 1
            continue

        if tok.type == "table_open":
            table_tokens = []
            j = i
            while j < len(tokens) and tokens[j].type != "table_close":
                table_tokens.append(tokens[j])
                j += 1
            table_tokens.append(tokens[j])  # table_close
            _render_table(doc, table_tokens, md_dir)
            i = j + 1
            continue

        if tok.type == "hr":
            p = doc.add_paragraph()
            ppr = p._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/>'
                f'</w:pBdr>'
            )
            ppr.append(pBdr)
            i += 1
            continue

        i += 1


def _find_matching_close(tokens: list[Token], start: int) -> int:
    open_type = tokens[start].type
    close_type = open_type.replace("_open", "_close")
    depth = 0
    for j in range(start, len(tokens)):
        if tokens[j].type == open_type:
            depth += 1
        elif tokens[j].type == close_type:
            depth -= 1
            if depth == 0:
                return j
    return len(tokens) - 1


def _render_table(doc: Document, tokens: list[Token], md_dir: Path) -> None:
    headers = []
    rows = []
    alignments = []
    current_row = []
    in_head = False
    in_body = False

    for tok in tokens:
        if tok.type == "thead_open":
            in_head = True
        elif tok.type == "thead_close":
            in_head = False
        elif tok.type == "tbody_open":
            in_body = True
        elif tok.type == "tbody_close":
            in_body = False
        elif tok.type == "tr_open":
            current_row = []
        elif tok.type == "tr_close":
            if in_head:
                headers = current_row
            elif in_body:
                rows.append(current_row)
        elif tok.type in ("th_open", "td_open"):
            style = tok.attrGet("style") or ""
            align = None
            if "text-align:center" in style:
                align = "center"
            elif "text-align:right" in style:
                align = "right"
            elif "text-align:left" in style:
                align = "left"
            if in_head:
                alignments.append(align)
        elif tok.type == "inline":
            current_row.append(tok.content)

    if headers:
        add_table(doc, headers, rows, alignments)


def embed_fonts(doc: Document) -> None:
    """Ensure Outfit fonts are embedded in the document."""
    # The template already has embedded fonts, but if we need to re-embed:
    regular_path = FONT_DIR / "Outfit-regular.ttf"
    bold_path = FONT_DIR / "Outfit-bold.ttf"

    if not regular_path.exists() or not bold_path.exists():
        return

    # Fonts are already in the template — no action needed
    # If the template is regenerated without fonts, this would need to
    # manipulate the docx ZIP to add them back


def convert(input_path: str, output_path: str | None = None) -> str:
    input_file = Path(input_path).resolve()
    if not input_file.exists():
        raise FileNotFoundError(f"Input file not found: {input_file}")

    if output_path is None:
        output_file = input_file.with_suffix(".docx")
    else:
        output_file = Path(output_path).resolve()

    md_text = input_file.read_text(encoding="utf-8")
    md_dir = input_file.parent

    doc = create_document()
    tokens = parse_markdown(md_text)
    render_tokens(doc, tokens, md_dir)

    doc.save(str(output_file))
    return str(output_file)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python md2docx.py input.md [output.docx]")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None

    result = convert(input_path, output_path)
    print(f"Created: {result}")
